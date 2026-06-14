from __future__ import annotations

import csv
import hashlib
import io
import logging
import math
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import httpx
import openpyxl

from app.core.config import get_settings

try:
    import pypdf
except ImportError:  # pragma: no cover - optional dependency guard
    pypdf = None

try:
    import docx
except ImportError:  # pragma: no cover - optional dependency guard
    docx = None

try:
    import xlrd
except ImportError:  # pragma: no cover - optional dependency guard
    xlrd = None


TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".json", ".log", ".yaml", ".yml", ".html", ".xml"}
EXCEL_EXTENSIONS = {".xlsx", ".xlsm", ".xls", ".csv"}
EMBEDDING_DIMENSIONS = 128
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 160
logger = logging.getLogger(__name__)


@dataclass
class ParsedContent:
    text: str
    parser: str
    metadata: dict[str, Any]


@dataclass
class TableData:
    headers: list[str]
    rows: list[list[str]]
    source_sheet: str = ""


def normalize_extension(name: str, extension: str | None = None) -> str:
    raw = extension or Path(name).suffix
    if not raw:
        return ""
    raw = raw.lower().strip()
    return raw if raw.startswith(".") else f".{raw}"


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse_file_bytes(filename: str, extension: str | None, data: bytes) -> ParsedContent:
    ext = normalize_extension(filename, extension)
    if ext == ".pdf":
        return _parse_pdf(data)
    if ext == ".docx":
        return _parse_docx(data)
    if ext in {".xlsx", ".xlsm"}:
        return _parse_xlsx(data)
    if ext == ".xls":
        return _parse_xls(data)
    if ext == ".csv":
        text = _csv_to_text(data)
        return ParsedContent(text=text, parser="csv", metadata={"extension": ext})
    if ext in TEXT_EXTENSIONS or not ext:
        return ParsedContent(text=decode_text(data), parser="text", metadata={"extension": ext})
    return ParsedContent(
        text=decode_text(data),
        parser="binary-text-fallback",
        metadata={"extension": ext, "warning": "Unsupported binary format, decoded as text where possible."},
    )


def _parse_pdf(data: bytes) -> ParsedContent:
    if pypdf is None:
        return ParsedContent(
            text="",
            parser="pdf-unavailable",
            metadata={"warning": "pypdf is not installed. Install dependencies to parse PDF files."},
        )
    reader = pypdf.PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {index}]\n{text.strip()}")
    return ParsedContent(text="\n\n".join(pages), parser="pdf", metadata={"pages": len(reader.pages)})


def _parse_docx(data: bytes) -> ParsedContent:
    if docx is None:
        return ParsedContent(
            text="",
            parser="docx-unavailable",
            metadata={"warning": "python-docx is not installed. Install dependencies to parse Word files."},
        )
    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append("\t".join(values))
    return ParsedContent(text="\n".join(parts), parser="docx", metadata={"paragraphs": len(document.paragraphs)})


def _parse_xlsx(data: bytes) -> ParsedContent:
    workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    sheets = []
    try:
        for sheet in workbook.worksheets:
            sheets.append(sheet.title)
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = [_cell_to_text(value) for value in row]
                if any(values):
                    parts.append("\t".join(values))
    finally:
        workbook.close()
    return ParsedContent(text="\n".join(parts), parser="xlsx", metadata={"sheets": sheets})


def _parse_xls(data: bytes) -> ParsedContent:
    if xlrd is None:
        return ParsedContent(
            text="",
            parser="xls-unavailable",
            metadata={"warning": "xlrd is not installed. Install dependencies to parse legacy Excel files."},
        )
    workbook = xlrd.open_workbook(file_contents=data)
    parts: list[str] = []
    for sheet in workbook.sheets():
        parts.append(f"[Sheet: {sheet.name}]")
        for row_idx in range(sheet.nrows):
            values = [_cell_to_text(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
            if any(values):
                parts.append("\t".join(values))
    return ParsedContent(text="\n".join(parts), parser="xls", metadata={"sheets": workbook.sheet_names()})


def _csv_to_text(data: bytes) -> str:
    table = parse_table_bytes("data.csv", ".csv", data)
    lines = []
    if table.headers:
        lines.append("\t".join(table.headers))
    lines.extend("\t".join(row) for row in table.rows)
    return "\n".join(lines)


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return text


def build_chunks(text: str, *, source_title: str = "", max_chars: int = CHUNK_SIZE) -> list[dict[str, Any]]:
    cleaned = re.sub(r"\n{3,}", "\n\n", text.replace("\r\n", "\n")).strip()
    if not cleaned:
        return []
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", cleaned) if part.strip()]
    chunks: list[dict[str, Any]] = []
    buffer = ""
    for paragraph in paragraphs:
        if len(buffer) + len(paragraph) + 2 <= max_chars:
            buffer = f"{buffer}\n\n{paragraph}".strip()
            continue
        if buffer:
            chunks.append(_chunk_payload(buffer, len(chunks), source_title))
        if len(paragraph) <= max_chars:
            buffer = paragraph
            continue
        for start in range(0, len(paragraph), max_chars - CHUNK_OVERLAP):
            part = paragraph[start : start + max_chars].strip()
            if part:
                chunks.append(_chunk_payload(part, len(chunks), source_title))
        buffer = ""
    if buffer:
        chunks.append(_chunk_payload(buffer, len(chunks), source_title))
    return chunks


def _chunk_payload(content: str, index: int, source_title: str) -> dict[str, Any]:
    return {
        "chunk_index": index,
        "content": content,
        "metadata": {"sourceTitle": source_title, "charCount": len(content)},
    }


def _normalize_vector(vector: list[float]) -> list[float]:
    norm = math.sqrt(sum(value * value for value in vector)) or 1.0
    return [round(value / norm, 6) for value in vector]


def _local_embed_text(text: str, dimensions: int) -> list[float]:
    tokens = re.findall(r"[\w\u4e00-\u9fff]+", text.lower())
    vector = [0.0] * dimensions
    if not tokens:
        return vector
    for token in tokens:
        digest = hashlib.sha256(token.encode("utf-8")).digest()
        index = int.from_bytes(digest[:4], "big") % dimensions
        weight = 1.0 + (int.from_bytes(digest[4:6], "big") % 100) / 1000
        vector[index] += weight
    return _normalize_vector(vector)


def embedding_model_name() -> str:
    settings = get_settings()
    if settings.embedding_api_key:
        return settings.embedding_model
    return f"local-hash-{settings.embedding_dimensions or EMBEDDING_DIMENSIONS}"


def _call_embedding_gateway(text: str, dimensions: int) -> list[float] | None:
    settings = get_settings()
    if not settings.embedding_api_key:
        return None

    base_url = settings.embedding_base_url.rstrip("/")
    payload: dict[str, Any] = {
        "model": settings.embedding_model,
        "input": text[:24000],
    }
    if dimensions and settings.embedding_model.startswith("text-embedding-3"):
        payload["dimensions"] = dimensions

    headers = {
        "Authorization": f"Bearer {settings.embedding_api_key}",
        "Content-Type": "application/json",
    }

    try:
        with httpx.Client(timeout=settings.ai_gateway_timeout) as client:
            response = client.post(f"{base_url}/embeddings", headers=headers, json=payload)
            if response.status_code >= 400 and "dimensions" in payload:
                payload.pop("dimensions", None)
                response = client.post(f"{base_url}/embeddings", headers=headers, json=payload)
            response.raise_for_status()
        data = response.json()
        vector = data["data"][0]["embedding"]
        if not isinstance(vector, list) or not vector:
            return None
        return _normalize_vector([float(value) for value in vector])
    except Exception as exc:  # pragma: no cover - depends on external provider
        logger.warning("Embedding gateway unavailable, falling back to local vectors: %s", exc)
        return None


def embed_text(text: str, dimensions: int | None = None) -> list[float]:
    settings = get_settings()
    size = dimensions or settings.embedding_dimensions or EMBEDDING_DIMENSIONS
    return _call_embedding_gateway(text, size) or _local_embed_text(text, size)


def cosine_similarity(left: list[float], right: list[float]) -> float:
    if not left or not right:
        return 0.0
    size = min(len(left), len(right))
    return sum(left[index] * right[index] for index in range(size))


def rank_chunks(question: str, chunks: list[dict[str, Any]], limit: int = 5) -> list[dict[str, Any]]:
    query_embedding = embed_text(question)
    ranked = []
    for item in chunks:
        embedding = item.get("embedding")
        if not isinstance(embedding, list):
            embedding = embed_text(str(item.get("content", "")))
        ranked.append({**item, "score": cosine_similarity(query_embedding, embedding)})
    ranked.sort(key=lambda item: item["score"], reverse=True)
    return ranked[:limit]


def _chat_with_gateway(messages: list[dict[str, str]], max_tokens: int = 1200, temperature: float = 0.2) -> str | None:
    settings = get_settings()
    if not settings.ai_gateway_api_key:
        return None

    payload: dict[str, Any] = {
        "model": settings.ai_gateway_model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    headers = {
        "Authorization": f"Bearer {settings.ai_gateway_api_key}",
        "Content-Type": "application/json",
    }
    try:
        with httpx.Client(timeout=settings.ai_gateway_timeout) as client:
            response = client.post(
                f"{settings.ai_gateway_base_url.rstrip('/')}/chat/completions",
                headers=headers,
                json=payload,
            )
            response.raise_for_status()
        data = response.json()
        content = data["choices"][0]["message"]["content"]
        return str(content).strip()
    except Exception as exc:  # pragma: no cover - depends on external provider
        logger.warning("AI gateway unavailable, falling back to local response: %s", exc)
        return None


def summarize_text(text: str, max_sentences: int = 5) -> str:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return "暂时没有可总结的文本内容。"
    ai_summary = _chat_with_gateway(
        [
            {
                "role": "system",
                "content": "你是 Lumio 序光的文件总结助手。请用中文输出结构化摘要，包含核心结论、关键数据、风险与下一步建议。",
            },
            {"role": "user", "content": cleaned[:12000]},
        ],
        max_tokens=1200,
        temperature=0.2,
    )
    if ai_summary:
        return ai_summary[:3000]
    sentences = re.split(r"(?<=[。！？.!?])\s+", cleaned)
    useful = [sentence.strip() for sentence in sentences if len(sentence.strip()) >= 12]
    if not useful:
        return cleaned[:600]
    selected = useful[:max_sentences]
    summary = "\n".join(f"{index + 1}. {sentence}" for index, sentence in enumerate(selected))
    return summary[:1800]


def answer_question(question: str, ranked_chunks: list[dict[str, Any]]) -> str:
    if not ranked_chunks:
        return "当前文件或知识库还没有可检索的内容，请先上传、解析或登记资料来源。"
    top = ranked_chunks[:3]
    evidence = "\n".join(
        f"- {str(item.get('content', '')).strip()[:260]}" for item in top if str(item.get("content", "")).strip()
    )
    context = "\n\n".join(
        f"来源：{item.get('title') or item.get('sourceType') or '资料'}\n{str(item.get('content', '')).strip()[:1800]}"
        for item in ranked_chunks[:6]
        if str(item.get("content", "")).strip()
    )
    ai_answer = _chat_with_gateway(
        [
            {
                "role": "system",
                "content": (
                    "你是 Lumio 序光的文件问答助手。只根据给定资料回答，资料不足时说明缺口。"
                    "回答要清晰、可执行，并在末尾列出引用来源标题。"
                ),
            },
            {"role": "user", "content": f"问题：{question}\n\n可用资料：\n{context[:10000]}"},
        ],
        max_tokens=1400,
        temperature=0.2,
    )
    if ai_answer:
        return ai_answer
    return (
        "根据已解析内容，优先参考以下片段：\n"
        f"{evidence}\n\n"
        f"针对你的问题“{question}”，建议先围绕上述片段核对原文；后续接入正式大模型后可生成更完整的推理答案。"
    )


def parse_table_bytes(filename: str, extension: str | None, data: bytes) -> TableData:
    ext = normalize_extension(filename, extension)
    if ext == ".csv":
        text = decode_text(data)
        rows = [[cell.strip() for cell in row] for row in csv.reader(io.StringIO(text))]
        return _rows_to_table(rows, source_sheet="CSV")
    if ext in {".xlsx", ".xlsm"}:
        workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        try:
            sheet = workbook.active
            rows = [[_cell_to_text(value) for value in row] for row in sheet.iter_rows(values_only=True)]
            return _rows_to_table(rows, source_sheet=sheet.title)
        finally:
            workbook.close()
    if ext == ".xls":
        if xlrd is None:
            raise ValueError("当前环境缺少 xlrd，无法解析 .xls 文件。")
        workbook = xlrd.open_workbook(file_contents=data)
        sheet = workbook.sheet_by_index(0)
        rows = [
            [_cell_to_text(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
            for row_idx in range(sheet.nrows)
        ]
        return _rows_to_table(rows, source_sheet=sheet.name)
    raise ValueError("当前文件不是可处理的表格格式。")


def _rows_to_table(rows: list[list[str]], source_sheet: str) -> TableData:
    rows = [row for row in rows if any(cell.strip() for cell in row)]
    if not rows:
        return TableData(headers=[], rows=[], source_sheet=source_sheet)
    headers = [cell.strip() or f"字段{index + 1}" for index, cell in enumerate(rows[0])]
    width = len(headers)
    normalized_rows = []
    for row in rows[1:]:
        padded = (row + [""] * width)[:width]
        normalized_rows.append([cell.strip() for cell in padded])
    return TableData(headers=headers, rows=normalized_rows, source_sheet=source_sheet)


def clean_table(table: TableData) -> TableData:
    seen: set[tuple[str, ...]] = set()
    cleaned: list[list[str]] = []
    for row in table.rows:
        normalized = [cell.strip() for cell in row]
        if not any(normalized):
            continue
        key = tuple(normalized)
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(normalized)
    return TableData(headers=[header.strip() for header in table.headers], rows=cleaned, source_sheet=table.source_sheet)


def split_table(table: TableData, *, column: str | None = None, rows_per_file: int | None = None) -> dict[str, TableData]:
    if column:
        if column not in table.headers:
            raise ValueError(f"未找到拆分列：{column}")
        index = table.headers.index(column)
        groups: dict[str, list[list[str]]] = {}
        for row in table.rows:
            key = row[index].strip() or "未分类"
            safe_key = re.sub(r'[\\/:*?"<>|]+', "_", key)[:80] or "未分类"
            groups.setdefault(safe_key, []).append(row)
        return {
            key: TableData(headers=table.headers, rows=rows, source_sheet=table.source_sheet)
            for key, rows in groups.items()
        }
    size = max(1, int(rows_per_file or 100))
    return {
        f"第{index + 1}批": TableData(headers=table.headers, rows=table.rows[start : start + size], source_sheet=table.source_sheet)
        for index, start in enumerate(range(0, len(table.rows), size))
    }


def merge_tables(tables: list[TableData]) -> TableData:
    if not tables:
        return TableData(headers=[], rows=[])
    headers = tables[0].headers
    merged: list[list[str]] = []
    for table in tables:
        if table.headers != headers:
            raise ValueError("文件表头不一致，无法直接合并。请先清洗或统一字段。")
        merged.extend(table.rows)
    return TableData(headers=headers, rows=merged)


def table_to_csv_bytes(table: TableData) -> bytes:
    output = io.StringIO()
    writer = csv.writer(output)
    if table.headers:
        writer.writerow(table.headers)
    writer.writerows(table.rows)
    return output.getvalue().encode("utf-8-sig")
